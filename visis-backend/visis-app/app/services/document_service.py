
# app/services/document_service.py

# import io
# import time
# import logging
# from datetime import datetime, timezone
# from concurrent.futures import ThreadPoolExecutor
# from app.core.config import settings
# from app.utils.redis_utils import redis_client, delete_pattern
# from pydub import AudioSegment
# from pydub.effects import normalize
# import json
# from sqlalchemy.orm import Session, Query
# from fastapi import HTTPException, status
# from typing import List, Optional
# from typing import Tuple

# from app.schemas.document import DocumentCreate, DocumentUpdate, DocumentFilter
# from app.models import Document, Audiobook,User 
# from app.models.user import SubscriptionType
# from app.schemas import AudioBookCreate
# from sqlalchemy import or_
# from app.utils.s3_utils import   s3_handler
# from app.utils.lang_utils  import map_language_code_to_supported, infer_genre
# from app.database import SessionLocal
# from app.services.audiobook_service import create_audiobook
# from app.services.ocr_service import OCRService
# from app.services.tts_service import TTSService
# from app.services.rekognition_service import RekognitionService
# from app.schemas.document_filter import DocumentFilter
# from docx import Document as DocxDocument

# logger = logging.getLogger(__name__)
# # Initialize services
# ocr_service = OCRService(region_name='us-east-1')
# tts_service = TTSService(region_name='us-east-1')
# rekognition_service = RekognitionService(region_name=settings.REGION_NAME)

# def add_delay(audio: AudioSegment, delay_ms: int = 50, attenuation_db: float = 10.0) -> AudioSegment:
#     """
#     Adds a delayed and attenuated copy of the audio to itself to simulate reverb.

#     :param audio: Original AudioSegment.
#     :param delay_ms: Delay in milliseconds.
#     :param attenuation_db: Attenuation in decibels for the delayed copy.
#     :return: AudioSegment with delay effect applied.
#     """
#     # Create a silent audio segment for the delay
#     silent_segment = AudioSegment.silent(duration=delay_ms)

#     # Create the delayed copy
#     delayed_copy = silent_segment + audio

#     # Attenuate the delayed copy
#     delayed_copy = delayed_copy - attenuation_db

#     # Overlay the delayed copy onto the original audio
#     return audio.overlay(delayed_copy)


# def apply_immersive_effects(audio: AudioSegment) -> AudioSegment:
#     """
#     Apply immersive audio effects including stereo widening and subtle reverb.
#     """
#     # Normalize audio volume
#     audio = normalize(audio)

#     # Stereo widening by slightly panning the audio
#     left = audio.pan(-0.3)   # Pan 30% to the left
#     right = audio.pan(0.3)   # Pan 30% to the right

#     # Combine left and right to create a wider stereo effect
#     immersive_audio = left.overlay(right)

#     # Add a subtle reverb effect by adding delay
#     immersive_audio = add_delay(immersive_audio, delay_ms=50, attenuation_db=10)

#     # Normalize again after effects
#     immersive_audio = normalize(immersive_audio)

#     return immersive_audio


# def process_document(document_id: int, user_id: int, bucket_name: str, file_key: str):
#     """
#     Background task to process the uploaded document.
#     """
#     db: Session = SessionLocal()
#     try:
#         document: Document = db.query(Document).filter(Document.id == document_id).first()
#         if not document:
#             logger.error(f"Document with ID {document_id} not found.")
#             return

#         user: User = db.query(User).filter(User.id == user_id).first()
#         if not user:
#             logger.error(f"User with ID {user_id} not found.")
#             raise Exception("User not found.")

#         subscription_type = user.subscription_type  # 'free' or 'premium'
#         is_dolby_atmos_supported = subscription_type == SubscriptionType.premium

#         logger.info(f"Starting processing for document ID: {document_id}")
#         start_time = datetime.utcnow()

#         # Extract text based on file type
#         if document.file_type == "application/pdf":
#             logger.info("Extracting text from PDF.")
#             text, detected_language = ocr_service.extract_text_from_pdf(bucket_name, file_key)
#         elif document.file_type.startswith("image/"):
#             logger.info("Extracting text from image.")
#             text, detected_language = ocr_service.extract_text_from_image(bucket_name, file_key)
#             logger.info("Detecting labels using Rekognition.")
#             tags = rekognition_service.detect_labels(bucket_name, file_key)
#             if isinstance(tags, str):
#                 tags = tags.split(", ")
#             if not tags:
#                 tags = ["No tags detected"]
#             document.tags = tags
#         elif document.file_type == "text/plain":
#             logger.info("Extracting text from TXT.")
#             file_content = s3_handler.get_file_content(bucket_name, file_key)
#             if not file_content:
#                 raise Exception("Failed to retrieve file content from S3.")
#             text = file_content.decode("utf-8")
#             detected_language = "en"
#         elif document.file_type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
#             logger.info("Extracting text from DOCX.")
#             file_content = s3_handler.get_file_content(bucket_name, file_key)
#             if not file_content:
#                 raise Exception("Failed to retrieve file content from S3.")
#             with io.BytesIO(file_content) as file_stream:
#                 docx = DocxDocument(file_stream)
#                 text = "\n".join(paragraph.text for paragraph in docx.paragraphs)

#             detected_language = ocr_service.detect_language(text)
#         else:
#             logger.error("Unsupported file type for processing.")
#             raise HTTPException(
#                 status_code=status.HTTP_400_BAD_REQUEST,
#                 detail="Unsupported file type for processing.",
#             )

#         if not text.strip():
#             logger.error("No text found in the document.")
#             raise Exception("No text found in the document.")

#         # Map detected language
#         full_language_name = map_language_code_to_supported(detected_language)
#         document.detected_language = full_language_name

#         # Generate description
#         document.description = text[:200].strip()

#         # Infer genre
#         genre = infer_genre(text)
#         document.genre = genre

#         # Generate audio
#         logger.info("Starting TTS conversion.")
#         tts_start_time = datetime.utcnow()

#         if len(text) <= 3000:
#             text_chunks = [text]
#             logger.info("Processing as a single text chunk.")
#         else:
#             max_length = 3000
#             text_chunks = [text[i:i + max_length] for i in range(0, len(text), max_length)]
#             logger.info(f"Processing {len(text_chunks)} text chunks.")

#         audio_segments = []

#         max_workers = min(12, len(text_chunks)) if len(text_chunks) > 1 else 1
#         logger.info(f"Using {max_workers} worker(s) for TTS conversion.")

#         with ThreadPoolExecutor(max_workers=max_workers) as executor:
#             audio_chunks = list(executor.map(
#                 lambda chunk: tts_service.convert_text_to_speech(chunk, detected_language),
#                 text_chunks,
#             ))

#         tts_end_time = datetime.utcnow()
#         tts_duration = (tts_end_time - tts_start_time).total_seconds()
#         logger.info(f"TTS conversion completed in {tts_duration} seconds.")

#         for i, chunk_audio in enumerate(audio_chunks):
#             if not chunk_audio:
#                 logger.warning(f"TTS conversion failed for chunk {i+1}.")
#                 continue
#             audio_segment = AudioSegment.from_mp3(io.BytesIO(chunk_audio))
#             audio_segments.append(audio_segment)

#         if not audio_segments:
#             logger.error("No audio segments were created from TTS.")
#             raise Exception("No audio segments were created from TTS.")

#         # Combine audio segments efficiently
#         logger.info("Combining audio segments.")
#         combine_start_time = datetime.utcnow()

#         combined_audio = sum(audio_segments)

#         # Apply immersive effects based on subscription_type
#         if is_dolby_atmos_supported:
#             immersive_audio = apply_immersive_effects(combined_audio)
#             logger.info("Applied immersive audio effects.")
#         else:
#             immersive_audio = combined_audio  # No immersive effects for free users
#             logger.info("No immersive audio effects applied for free user.")

#         combine_end_time = datetime.utcnow()
#         combine_duration = (combine_end_time - combine_start_time).total_seconds()
#         logger.info(f"Audio combining and effects applied in {combine_duration} seconds.")

#         # Calculate duration
#         duration_in_seconds = len(immersive_audio) / 1000
#         duration_in_minutes = duration_in_seconds // 60
#         duration_in_hours = duration_in_minutes // 60
#         formatted_duration = f"{int(duration_in_hours)} hours {int(duration_in_minutes % 60)} minutes"

#         # Export audio
#         logger.info("Exporting audio to BytesIO.")
#         export_start_time = datetime.utcnow()

#         audio_bytes_io = io.BytesIO()
#         immersive_audio.export(audio_bytes_io, format="mp3", bitrate="128k")
#         audio_bytes_io.seek(0)

#         export_end_time = datetime.utcnow()
#         export_duration = (export_end_time - export_start_time).total_seconds()
#         logger.info(f"Audio export completed in {export_duration} seconds.")

#         # Upload audio
#         logger.info("Uploading processed audio to S3.")
#         upload_start_time = datetime.utcnow()

#         processed_audio_key = f"audio/{user_id}/{document.id}_processed.mp3"
#         audio_url = s3_handler.upload_file(
#             file_obj=audio_bytes_io,
#             bucket=bucket_name,
#             key=processed_audio_key,
#             content_type="audio/mpeg",
#         )
#         if not audio_url:
#             logger.error("Failed to upload processed audio to S3.")
#             raise Exception("Failed to upload processed audio to S3.")

#         upload_end_time = datetime.utcnow()
#         upload_duration = (upload_end_time - upload_start_time).total_seconds()
#         logger.info(f"Uploaded processed audio to S3 at {audio_url} in {upload_duration} seconds.")

#         # Create audiobook
#         audiobook_data = AudioBookCreate(
#             title=document.title,
#             author=document.author,
#             narrator="Generated Narrator",
#             duration=formatted_duration,
#             genre=genre,
#             publication_date=datetime.utcnow(),
#             file_key=processed_audio_key,
#             is_dolby_atmos_supported=is_dolby_atmos_supported,
#             url=audio_url,
#             document_id=document.id,
#         )
#         audiobook = create_audiobook(db=db, audiobook_data=audiobook_data)

#         # Link audiobook to document
#         document.audiobook = audiobook
#         document.audio_url = audio_url  # Ensure audio_url is updated
#         document.processing_status = "completed"
#         db.commit()

#         end_time = datetime.utcnow()
#         total_processing_time = (end_time - start_time).total_seconds()
#         logger.info(f"Total processing time for document {document.id}: {total_processing_time:.2f} seconds.")

#         # Optionally, enforce processing time constraints
#         if total_processing_time > 60:
#             logger.warning(f"Processing time exceeded 60 seconds for document {document.id}: {total_processing_time} seconds.")

#     except HTTPException as http_exc:
#         logger.error(f"HTTP error processing document {document_id}: {http_exc.detail}")
#         if 'document' in locals():
#             try:
#                 document.processing_status = "failed"
#                 document.processing_error = str(http_exc.detail)
#                 db.commit()
#                 logger.info(f"Updated document {document_id} as failed due to HTTP error.")
#             except Exception as update_exc:
#                 db.rollback()
#                 logger.error(f"Failed to update document {document_id} after HTTP error: {str(update_exc)}")
#         raise http_exc

#     except Exception as e:
#         logger.error(f"Error processing document {document_id}: {str(e)}")
#         if 'document' in locals():
#             try:
#                 document.processing_status = "failed"
#                 document.processing_error = str(e)
#                 db.commit()
#                 logger.info(f"Updated document {document_id} as failed due to error.")
#             except Exception as update_exc:
#                 db.rollback()
#                 logger.error(f"Failed to update document {document_id} after error: {str(update_exc)}")
#         raise HTTPException(
#             status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
#             detail=f"Document processing failed: {str(e)}",
#         )
#     finally:
#         db.close()



# # def process_document(document_id: int, user_id: int, bucket_name: str, file_key: str):
# #     """
# #     Background task to process the uploaded document.
# #     """
# #     db: Session = SessionLocal()  # Create a new session
# #     try:
# #         # Retrieve the document from the database
# #         document: Document = db.query(Document).filter(Document.id == document_id).first()
# #         if not document:
# #             logger.error(f"Document with ID {document_id} not found.")
# #             return

        
        
# #         # Retrieve the user to get subscription_type
# #         user: User = db.query(User).filter(User.id == user_id).first()
# #         if not user:
# #             logger.error(f"User with ID {user_id} not found.")
# #             raise Exception("User not found.")
        
# #         subscription_type = user.subscription_type  # 'free' or 'premium'
# #         is_dolby_atmos_supported = subscription_type == SubscriptionType.premium

# #         logger.info(f"Starting processing for document ID: {document_id}")
# #         # Record the start time
# #         start_time = datetime.utcnow()

# #         # Extract text based on file type
# #         if document.file_type == "application/pdf":
# #             logger.info("Extracting text from PDF.")
# #             text, detected_language = ocr_service.extract_text_from_pdf(bucket_name, file_key)
# #         elif document.file_type.startswith("image/"):
# #             logger.info("Extracting text from image.")
# #             text, detected_language = ocr_service.extract_text_from_image(bucket_name, file_key)
# #             # Rekognition for tags
# #             logger.info("Detecting labels using Rekognition.")
# #             tags = rekognition_service.detect_labels(bucket_name, file_key)
# #             if isinstance(tags, str):
# #                 tags = tags.split(", ")
# #             if not tags:
# #                 tags = ["No tags detected"]
# #             document.tags = tags
# #         elif document.file_type == "text/plain":
# #             logger.info("Extracting text from TXT.")
# #             file_content = s3_handler.get_file_content(bucket_name, file_key)
# #             if not file_content:
# #                 raise Exception("Failed to retrieve file content from S3.")
# #             text = file_content.decode("utf-8")
# #             detected_language = "en"

# #         elif document.file_type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
# #             logger.info("Extracting text from DOCX.")
# #             file_content = s3_handler.get_file_content(bucket_name, file_key)
# #             if not file_content:
# #                 raise Exception("Failed to retrieve file content from S3.")
# #             with io.BytesIO(file_content) as file_stream:
# #                 docx = DocxDocument(file_stream)
# #                 text = "\n".join(paragraph.text for paragraph in docx.paragraphs)

# #             # Use OCRService to detect language automatically
# #             detected_language = ocr_service.detect_language(text)

# #         else:
# #             logger.error("Unsupported file type for processing.")
# #             raise HTTPException(
# #                 status_code=status.HTTP_400_BAD_REQUEST,
# #                 detail="Unsupported file type for processing.",
# #             )

# #         if not text.strip():
# #             logger.error("No text found in the document.")
# #             raise Exception("No text found in the document.")

# #         # Map detected language
# #         full_language_name = map_language_code_to_supported(detected_language)
# #         document.detected_language = full_language_name

# #         # Generate description
# #         document.description = text[:200].strip()

# #         # Infer genre
# #         genre = infer_genre(text)
# #         document.genre = genre

# #         # Generate audio
# #         logger.info("Starting TTS conversion.")
# #         tts_start_time = datetime.utcnow()

# #         # Optimize chunking: process as single chunk if possible
# #         if len(text) <= 3000:
# #             text_chunks = [text]
# #             logger.info("Processing as a single text chunk.")
# #         else:
# #             max_length = 3000
# #             text_chunks = [text[i:i + max_length] for i in range(0, len(text), max_length)]
# #             logger.info(f"Processing {len(text_chunks)} text chunks.")

# #         audio_segments = []

# #         # Adjust max_workers based on number of chunks
# #         max_workers = min(12, len(text_chunks)) if len(text_chunks) > 1 else 1
# #         logger.info(f"Using {max_workers} worker(s) for TTS conversion.")

# #         with ThreadPoolExecutor(max_workers=max_workers) as executor:
# #             audio_chunks = list(
# #                 executor.map(
# #                     lambda chunk: tts_service.convert_text_to_speech(chunk, detected_language),
# #                     text_chunks,
# #                 )
# #             )

# #         tts_end_time = datetime.utcnow()
# #         tts_duration = (tts_end_time - tts_start_time).total_seconds()
# #         logger.info(f"TTS conversion completed in {tts_duration} seconds.")

# #         for i, chunk_audio in enumerate(audio_chunks):
# #             if not chunk_audio:
# #                 logger.warning(f"TTS conversion failed for chunk {i+1}.")
# #                 continue
# #             audio_segment = AudioSegment.from_file(io.BytesIO(chunk_audio), format="mp3")
# #             audio_segments.append(audio_segment)

# #         if not audio_segments:
# #             logger.error("No audio segments were created from TTS.")
# #             raise Exception("No audio segments were created from TTS.")

# #         # Combine audio segments
# #         logger.info("Combining audio segments.")
# #         combine_start_time = datetime.utcnow()

# #         combined_audio = AudioSegment.empty()
# #         for segment in audio_segments:
# #             combined_audio += segment

# #         # Apply immersive effects based on subscription_type
# #         if is_dolby_atmos_supported:
# #             immersive_audio = apply_immersive_effects(combined_audio)
# #             logger.info("Applied immersive audio effects.")
# #         else:
# #             immersive_audio = combined_audio  # No immersive effects for free users
# #             logger.info("No immersive audio effects applied for free user.")


# #         # immersive_audio = apply_immersive_effects(combined_audio)

# #         combine_end_time = datetime.utcnow()
# #         combine_duration = (combine_end_time - combine_start_time).total_seconds()
# #         logger.info(f"Audio combining and effects applied in {combine_duration} seconds.")

# #         # Calculate duration
# #         duration_in_seconds = len(immersive_audio) / 1000
# #         duration_in_minutes = duration_in_seconds // 60
# #         duration_in_hours = duration_in_minutes // 60
# #         formatted_duration = f"{int(duration_in_hours)} hours {int(duration_in_minutes % 60)} minutes"

# #         # Export audio
# #         logger.info("Exporting audio to BytesIO.")
# #         export_start_time = datetime.utcnow()

# #         audio_bytes_io = io.BytesIO()
# #         immersive_audio.export(audio_bytes_io, format="mp3", bitrate="128k")
# #         audio_bytes_io.seek(0)

# #         export_end_time = datetime.utcnow()
# #         export_duration = (export_end_time - export_start_time).total_seconds()
# #         logger.info(f"Audio export completed in {export_duration} seconds.")

# #         # Upload audio
# #         logger.info("Uploading processed audio to S3.")
# #         upload_start_time = datetime.utcnow()

# #         processed_audio_key = f"audio/{user_id}/{document.id}_processed.mp3"
# #         audio_url = s3_handler.upload_file(
# #             file_obj=audio_bytes_io,
# #             bucket=bucket_name,
# #             key=processed_audio_key,
# #             content_type="audio/mpeg",
# #         )
# #         if not audio_url:
# #             logger.error("Failed to upload processed audio to S3.")
# #             raise Exception("Failed to upload processed audio to S3.")

# #         upload_end_time = datetime.utcnow()
# #         upload_duration = (upload_end_time - upload_start_time).total_seconds()
# #         logger.info(f"Uploaded processed audio to S3 at {audio_url} in {upload_duration} seconds.")

# #         # Create audiobook
# #         audiobook_data = AudioBookCreate(
# #             title=document.title,
# #             narrator="Generated Narrator",
# #             duration=formatted_duration,
# #             genre=genre,
# #             publication_date=datetime.utcnow(),
# #             author=document.author,
# #             file_key=processed_audio_key,
# #             url=audio_url,
# #             is_dolby_atmos_supported=is_dolby_atmos_supported,
# #             document_id=document.id,
# #         )
# #         audiobook = create_audiobook(db=db, audiobook=audiobook_data)

# #         # Link audiobook to document
# #         document.audiobook = audiobook
# #         document.processing_status = "completed"
# #         db.commit()

# #         # Calculate total processing time
# #         end_time = datetime.utcnow()
# #         total_processing_time = (end_time - start_time).total_seconds()
# #         logger.info(f"Total processing time for document {document.id}: {total_processing_time} seconds.")

# #         # Ensure processing time is under 60 seconds
# #         if total_processing_time > 60:
# #             logger.warning(f"Processing time exceeded 60 seconds for document {document.id}: {total_processing_time} seconds.")

# #     except HTTPException as http_exc:
# #         logger.error(f"HTTP error processing document {document_id}: {http_exc.detail}")
# #         if 'document' in locals():
# #             document.processing_status = "failed"
# #             document.processing_error = str(http_exc.detail)
# #             db.commit()
# #         raise http_exc
# #     except Exception as e:
# #         logger.error(f"Error processing document {document_id}: {str(e)}")
# #         if 'document' in locals():
# #             document.processing_status = "failed"
# #             document.processing_error = str(e)
# #             db.commit()
# #         raise HTTPException(
# #             status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
# #             detail=f"Document processing failed: {str(e)}",
# #         )
# #     finally:
# #         db.close()




# def create_document(
#     db: Session,
#     document_data: DocumentCreate,
#     user_id: int,
#     file_key: str,
#     file_size: Optional[int] = None,
# ) -> Document:
#     """Create a new document record and enqueue processing."""
#     try:
#         # Initialize the document with processing_status as 'processing'
#         document = Document(
#             title=document_data.title,
#             author=document_data.author,
#             file_type=document_data.file_type,
#             file_key=file_key,
#             url=document_data.url,
#             is_public=document_data.is_public,
#             user_id=user_id,
#             file_size=file_size,
#             upload_date=datetime.utcnow(),
#             created_at=document_data.created_at,
#             processing_status='processing'
#         )
#         db.add(document)
#         db.commit()
#         db.refresh(document)

#         logger.info(f"Created document record with ID {document.id}.")

#         return document
#     except Exception as e:
#         db.rollback()
#         logger.error(f"Error creating document: {str(e)}")
#         raise HTTPException(
#             status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
#             detail=f"Failed to create document: {str(e)}"
#         )



# # # ##working
# # # # ##working
# # def create_document(
# #     db: Session,
# #     document_data: DocumentCreate,
# #     user_id: int,
# #     file_key: str,
# #     file_size: Optional[int] = None,
# #     subscription_type: str = "free" 
# # ) -> Document:
# #     """Create a new document record with synchronous OCR and TTS processing."""
# #     try:
# #         # Initialize the document with processing_status as 'processing'
# #         document = Document(
# #             title=document_data.title,
# #             author=document_data.author,
# #             file_type=document_data.file_type,
# #             file_key=file_key,
# #             url=document_data.url,
# #             is_public=document_data.is_public,
# #             user_id=user_id,
# #             file_size=file_size,
# #             upload_date=datetime.utcnow(),
# #             created_at=document_data.created_at,
# #             processing_status='processing'
# #         )
# #         db.add(document)
# #         db.commit()
# #         db.refresh(document)

# #         logger.info(f"Created document record with ID {document.id}.")

# #         # Extract text based on file type
# #         if document.file_type == 'application/pdf':
# #             text, detected_language = ocr_service.extract_text_from_pdf(
# #                 bucket_name=settings.S3_BUCKET_NAME,
# #                 file_key=document.file_key
# #             )
# #         elif document.file_type.startswith('image/'):
# #             text, detected_language = ocr_service.extract_text_from_image(
# #                 bucket_name=settings.S3_BUCKET_NAME,
# #                 file_key=document.file_key
# #             )
# #         elif document.file_type == "text/plain":
# #             text, detected_language = ocr_service.extract_text_from_txt(
# #                 bucket_name=settings.S3_BUCKET_NAME,
# #                 file_key=document.file_key
# #             )
# #         else:
# #             raise Exception("Unsupported file type")

# #         if not text.strip():
# #             raise Exception("No text found in the document")

# #         logger.info(f"Extracted text from document ID {document.id}.")

# #         # Generate audio and upload to S3
# #         audio_url, processed_audio_key,is_dolby_atmos_supported = generate_audio_for_document(document, text, detected_language,subscription_type=subscription_type)

# #         # Language, description, genre
# #         full_language_name = map_language_code_to_supported(detected_language)
# #         document.detected_language = full_language_name
# #         document.description = text[:200].strip() if text.strip() else None
# #         inferred_genre = infer_genre(text)
# #         document.genre = inferred_genre if inferred_genre else "Unknown"

# #         # Update document with audio URL and processing status
# #         document.audio_url = audio_url
# #         document.audio_key = processed_audio_key
# #         document.processing_status = 'completed'
# #         document.detected_language = detected_language
# #         db.commit()
# #         db.refresh(document)

# #         logger.info(f"Completed processing for document ID {document.id}.")

# #         return document
# #     except Exception as e:
# #         db.rollback()
# #         logger.error(f"Error creating document: {str(e)}")
# #         raise HTTPException(
# #             status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
# #             detail=f"Failed to create document: {str(e)}"
# #         )




# def get_documents(
#     db: Session,
#     user_id: int,
#     filter_params: DocumentFilter,
#     skip: int = 0,
#     limit: int = 10
# ) -> List[Document]:
#     """
#     Retrieve documents belonging to a user with optional filters.
#     """
#     query = db.query(Document).filter(Document.user_id == user_id)
    
#     # Apply search filter
#     if filter_params.search:
#         search_term = f"%{filter_params.search}%"
#         query = query.filter(
#             or_(
#                 Document.title.ilike(search_term),
#                 Document.description.ilike(search_term)
#             )
#         )
    
#     # Apply status filter
#     if filter_params.status:
#         query = query.filter(Document.processing_status == filter_params.status)
    
#     # Apply file type filter
#     if filter_params.file_type:
#         query = query.filter(Document.file_type == filter_params.file_type)
    
#     # Apply date range filter
#     if filter_params.start_date:
#         query = query.filter(Document.created_at >= filter_params.start_date)
#     if filter_params.end_date:
#         query = query.filter(Document.created_at <= filter_params.end_date)
    
#     # Apply pagination
#     documents = query.order_by(Document.created_at.desc()).offset(skip).limit(limit).all()
    
#     return documents

# def count_user_pdfs(db: Session, user_id: int) -> int:
#     # Count only PDFs
#     return db.query(Document).filter(
#         Document.user_id == user_id,
#         Document.file_type == "application/pdf"
#     ).count()

# def update_document(
#     db: Session,
#     document_id: int,
#     user_id: int,
#     document_update: DocumentUpdate
# ) -> Document:
#     """Update an existing document record with Redis cache invalidation."""
#     try:
#         document = db.query(Document).filter(
#             Document.id == document_id,
#             Document.user_id == user_id
#         ).first()

#         if not document:
#             raise HTTPException(
#                 status_code=status.HTTP_404_NOT_FOUND,
#                 detail="Document not found"
#             )

#         # Update the document fields
#         for key, value in document_update.dict(exclude_unset=True).items():
#             setattr(document, key, value)

#         db.commit()
#         db.refresh(document)

#         # Invalidate relevant Redis cache
#         # delete_pattern("search:*")

#         return document
#     except Exception as e:
#         db.rollback()
#         logger.error(f"Error updating document {document_id}: {str(e)}")
#         raise HTTPException(
#             status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
#             detail="Failed to update document"
#         )


# def delete_document(db: Session, document_id: int, user_id: int) -> bool:
#     """Delete document and associated files with Redis cache invalidation."""
#     try:
#         document = db.query(Document).filter(
#             Document.id == document_id,
#             Document.user_id == user_id
#         ).first()

#         if not document:
#             raise HTTPException(
#                 status_code=status.HTTP_404_NOT_FOUND,
#                 detail="Document not found"
#             )

#         # Delete files from S3
#         if document.file_key:
#             s3_handler.delete_file(settings.S3_BUCKET_NAME, document.file_key)
#         if document.audio_key:
#             s3_handler.delete_file(settings.S3_BUCKET_NAME, document.audio_key)

#         # Delete database record
#         db.delete(document)
#         db.commit()

#         # Invalidate relevant Redis cache
#         # delete_pattern("search:*")

#         return True
#     except Exception as e:
#         db.rollback()
#         logger.error(f"Error deleting document {document_id}: {str(e)}")
#         raise HTTPException(
#             status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
#             detail="Failed to delete document"
#         )


# def search_documents_with_cache(db: Session, prefix: str, bucketname: str, limits: int, levels: int, offsets: int, user_id: int):
#     """Search documents with Redis caching."""
#     cache_key = f"search:{prefix}:{bucketname}:{limits}:{levels}:{offsets}"
#     cached_result = redis_client.get(cache_key)

#     if cached_result:
#         return json.loads(cached_result)

#     query = """
#     WITH files_folders AS (
#         SELECT ((string_to_array(objects.name, '/'))[%(levels)s]) AS folder
#         FROM storage.objects
#         WHERE objects.name ILIKE %(prefix)s || '%'
#         AND bucket_id = %(bucketname)s
#         AND user_id = %(user_id)s
#         GROUP BY folder
#         LIMIT %(limits)s
#         OFFSET %(offsets)s
#     )
#     SELECT files_folders.folder AS name, objects.id, objects.updated_at,
#            objects.created_at, objects.last_accessed_at, objects.metadata
#     FROM files_folders
#     LEFT JOIN storage.objects
#     ON %(prefix)s || files_folders.folder = objects.name
#        AND objects.bucket_id = %(bucketname)s
#        AND objects.user_id = %(user_id)s;
#     """
#     result = db.execute(query, {
#         "prefix": prefix,
#         "bucketname": bucketname,
#         "user_id": user_id,
#         "limits": limits,
#         "levels": levels,
#         "offsets": offsets,
#     }).fetchall()

#     serialized_result = [
#         {
#             "name": row.name,
#             "id": row.id,
#             "updated_at": row.updated_at,
#             "created_at": row.created_at,
#             "last_accessed_at": row.last_accessed_at,
#             "metadata": row.metadata,
#         }
#         for row in result
#     ]

#     redis_client.set(cache_key, json.dumps(serialized_result), ex=80640)
#     return serialized_result


# def process_text_chunks_in_parallel(text_chunks, detected_language):
#     """Process text chunks for TTS in parallel."""
#     max_workers = min(12, len(text_chunks))  # Optimize the number of workers
#     with ThreadPoolExecutor(max_workers=max_workers) as executor:
#         audio_chunks = list(executor.map(
#             lambda chunk: tts_service.convert_text_to_speech(chunk, detected_language),
#             text_chunks
#         ))
#     return audio_chunks


# # ##Working
# # from typing import Tuple
# def generate_audio_for_document(
#     document: Document,
#     text: str,
#     detected_language: str,
#     subscription_type: str = "free"  # Default to "free"
# ) -> Tuple[str, str,bool]:
#     """Generate audio for a document and upload to S3."""
#     try:
#         # Split text into manageable chunks (optimized size)
#         max_length = 3000  # Increase chunk size for fewer chunks
#         text_chunks = [text[i:i + max_length] for i in range(0, len(text), max_length)]

#         logger.info(f"Splitting text into {len(text_chunks)} chunks for TTS processing.")

#         # Process chunks in parallel
#         audio_chunks = process_text_chunks_in_parallel(text_chunks, detected_language)

#         # Combine audio chunks into one audio file
#         combined_audio = AudioSegment.empty()
#         for chunk_audio in audio_chunks:
#             combined_audio += AudioSegment.from_mp3(io.BytesIO(chunk_audio))

#         logger.info("Combining audio chunks into a single audio file.")

#         # Apply immersive effects only if premium
#         if subscription_type == "premium":
#             immersive_audio = apply_immersive_effects(combined_audio)
#             is_dolby_atmos_supported = True
#         else:
#             immersive_audio = combined_audio  # No immersive effects for free
#             is_dolby_atmos_supported = False

#         # Export processed audio to a BytesIO object
#         processed_audio_io = io.BytesIO()
#         immersive_audio.export(processed_audio_io, format="mp3", bitrate="128k")
#         processed_audio_io.seek(0)

#         logger.info("Exported audio to BytesIO object.")

#         # Upload processed audio to S3
#         processed_audio_key = f"{settings.S3_FOLDER_NAME}/audio/{document.user_id}/{document.id}_processed.mp3"
#         audio_url = s3_handler.upload_file(
#             file_obj=processed_audio_io,
#             bucket=settings.S3_BUCKET_NAME,
#             key=processed_audio_key,
#             content_type="audio/mpeg",
#         )
#         if not audio_url:
#             raise Exception("Failed to upload audio to S3")

#         logger.info(f"Uploaded processed audio to S3 at {audio_url}.")

#         return audio_url, processed_audio_key, is_dolby_atmos_supported

#     except Exception as e:
#         logger.error(f"Error generating audio for document {document.id}: {str(e)}")
#         raise e



# # app/services/document_service.py
# #V2

# import io
# import logging
# from datetime import datetime
# from concurrent.futures import ThreadPoolExecutor
# from typing import List, Optional

# from fastapi import HTTPException, status
# from sqlalchemy.orm import Session
# from sqlalchemy import or_

# from pydub import AudioSegment
# from pydub.effects import normalize
# from docx import Document as DocxDocument
# import time
# from app.core.config import settings
# from app.database import SessionLocal
# from app.models import Document, Audiobook, User
# from app.models.user import SubscriptionType
# from app.schemas.document import DocumentCreate, DocumentFilter, DocumentUpdate
# from app.schemas import AudioBookCreate
# from app.services.audiobook_service import create_audiobook
# from app.services.ocr_service import OCRService
# from app.services.tts_service import TTSService
# from app.services.rekognition_service import RekognitionService
# from app.utils.s3_utils import s3_handler
# from app.utils.lang_utils import map_language_code_to_supported, infer_genre
# # from app.utils.redis_utils import delete_pattern, redis_client  # Uncomment if you use Redis caching in search

# logger = logging.getLogger(__name__)

# # Initialize external services
# ocr_service = OCRService(region_name='us-east-1')
# tts_service = TTSService(region_name='us-east-1')
# rekognition_service = RekognitionService(region_name=settings.REGION_NAME)

# # ------------------------------------------------------------------------------
# #  IMMERSIVE AUDIO EFFECTS
# # ------------------------------------------------------------------------------
# def apply_immersive_effects(audio: AudioSegment) -> AudioSegment:
#     """Apply immersive audio effects including stereo widening and subtle reverb."""
#     # Normalize audio volume
#     audio = normalize(audio)

#     # Stereo widening
#     left = audio.pan(-0.3)
#     right = audio.pan(0.3)
#     immersive_audio = left.overlay(right)

#     # Subtle reverb/delay
#     delay_ms = 50
#     attenuation_db = 10.0
#     silent_segment = AudioSegment.silent(duration=delay_ms)
#     delayed_copy = (silent_segment + immersive_audio) - attenuation_db
#     immersive_audio = immersive_audio.overlay(delayed_copy)

#     # Normalize again
#     immersive_audio = normalize(immersive_audio)
#     return immersive_audio


# # ------------------------------------------------------------------------------
# #  CREATE & PROCESS DOCUMENT (SYNCHRONOUS)
# # ------------------------------------------------------------------------------
# def create_and_process_document(
#     db: Session,
#     document_data: DocumentCreate,
#     user: User,
#     file_content: bytes,
#     subscription_type: str = "free"
# ) -> Document:
#     """
#     Create a new Document, perform OCR + TTS + S3 uploads **synchronously**,
#     and attach the audiobook to the Document record. Return the final Document.
#     """

#     try:
#         # Start total processing time
#         total_start_time = time.time()
#         # 1) Insert Document row in DB (initially "processing")
#         document = Document(
#             title=document_data.title,
#             author=document_data.author,
#             file_type=document_data.file_type,
#             file_key=document_data.file_key,
#             url=document_data.url,
#             is_public=document_data.is_public,
#             user_id=user.id,
#             file_size=len(file_content),
#             upload_date=datetime.utcnow(),
#             created_at=document_data.created_at,
#             processing_status='processing'
#         )
#         db.add(document)
#         db.commit()
#         db.refresh(document)

#         logger.info(f"[SYNC] Document record {document.id} created. Starting processing...")

#         # 2) OCR / Text Extraction
#         ocr_start_time = time.time()
#         if document.file_type == "application/pdf":
#             text, detected_lang = ocr_service.extract_text_from_pdf(
#                 bucket_name=settings.S3_BUCKET_NAME,
#                 file_key=document.file_key
#             )

#         elif document.file_type.startswith("image/"):
#             text, detected_lang = ocr_service.extract_text_from_image(
#                 bucket_name=settings.S3_BUCKET_NAME,
#                 file_key=document.file_key
#             )
#             # Rekognition tags for images
#             tags_start_time = time.time()
#             tags = rekognition_service.detect_labels(settings.S3_BUCKET_NAME, document.file_key)
#             if isinstance(tags, str):
#                 tags = tags.split(", ")
#             document.tags = tags or ["No tags detected"]
#             tags_end_time = time.time()
#             logger.info(f"[SYNC] Rekognition tags for Document {document.id} processed in {tags_end_time - tags_start_time:.2f} seconds.")

#         elif document.file_type == "text/plain":
#             # Read text directly from S3
#             file_bytes = s3_handler.get_file_content(settings.S3_BUCKET_NAME, document.file_key)
#             text = file_bytes.decode("utf-8")
#             detected_lang = ocr_service.detect_language(text)

#         elif document.file_type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
#             # Handle DOCX
#             file_bytes = s3_handler.get_file_content(settings.S3_BUCKET_NAME, document.file_key)
#             with io.BytesIO(file_bytes) as docx_stream:
#                 docx_obj = DocxDocument(docx_stream)
#                 text = "\n".join(p.text for p in docx_obj.paragraphs)
#             detected_lang = ocr_service.detect_language(text)

#         else:
#             msg = f"Unsupported file type: {document.file_type}"
#             logger.error(msg)
#             raise HTTPException(status_code=400, detail=msg)

#         if not text.strip():
#             msg = f"No text found in document {document.id}"
#             logger.error(msg)
#             raise HTTPException(status_code=400, detail=msg)
        
#         ocr_end_time = time.time()
#         logger.info(f"[SYNC] OCR for Document {document.id} completed in {ocr_end_time - ocr_start_time:.2f} seconds.")

#         # 3) Post-OCR updates
#         post_ocr_start_time = time.time()
#         document.detected_language = map_language_code_to_supported(detected_lang or "en")
#         document.description = text[:200].strip()  # First 200 chars
#         document.genre = infer_genre(text)
#         db.commit()
#         db.refresh(document)
#         post_ocr_end_time = time.time()
#         logger.info(f"[SYNC] Post-OCR updates for Document {document.id} completed in {post_ocr_end_time - post_ocr_start_time:.2f} seconds.")

#         # 4) TTS Generation
#         tts_start_time = time.time()
#         is_dolby_atmos = (subscription_type == "premium")
#         logger.info(f"[SYNC] Generating audio for Document {document.id} (Dolby={is_dolby_atmos})")

#         # Split large text into chunks of ~3000 chars
#         max_chars = 3000
#         if len(text) <= max_chars:
#             text_chunks = [text]
#         else:
#             text_chunks = [text[i : i + max_chars] for i in range(0, len(text), max_chars)]

#         audio_segments = []
#         # TTS in parallel (up to 12 workers, or len(text_chunks) if fewer)
#         max_workers = min(12, len(text_chunks))
#         with ThreadPoolExecutor(max_workers=max_workers) as executor:
#             polly_start_time = time.time()
#             polly_results = list(
#                 executor.map(
#                     lambda chunk: tts_service.convert_text_to_speech(chunk, detected_lang or "en"),
#                     text_chunks
#                 )
#             )
#             polly_end_time = time.time()
#             logger.info(f"[SYNC] TTS conversion for Document {document.id} completed in {polly_end_time - polly_start_time:.2f} seconds.")

#         # Combine all segments
#         for chunk_data in polly_results:
#             if chunk_data:
#                 audio_segment = AudioSegment.from_mp3(io.BytesIO(chunk_data))
#                 audio_segments.append(audio_segment)

#         if not audio_segments:
#             msg = f"TTS failed: no audio segments produced for document {document.id}"
#             logger.error(msg)
#             raise HTTPException(status_code=500, detail=msg)

#         combined_audio = sum(audio_segments)

#         # Apply immersive effects for premium users
#         if is_dolby_atmos:
#             immersive_start_time = time.time()
#             combined_audio = apply_immersive_effects(combined_audio)
#             immersive_end_time = time.time()
#             logger.info(f"[SYNC] Applied immersive effects to Document {document.id} in {immersive_end_time - immersive_start_time:.2f} seconds.")

#         # 5) Export final MP3
#         export_start_time = time.time()
#         final_audio_bytes = io.BytesIO()
#         combined_audio.export(final_audio_bytes, format="mp3", bitrate="128k")
#         final_audio_bytes.seek(0)
#         export_end_time = time.time()
#         logger.info(f"[SYNC] Exported audio for Document {document.id} in {export_end_time - export_start_time:.2f} seconds.")

#         # 6) Upload processed MP3 to S3
#         upload_start_time = time.time()
#         processed_audio_key = f"audio/{user.id}/{document.id}_processed.mp3"
#         audio_url = s3_handler.upload_file(
#             file_obj=final_audio_bytes,
#             bucket=settings.S3_BUCKET_NAME,
#             key=processed_audio_key,
#             content_type="audio/mpeg"
#         )
#         upload_end_time = time.time()
#         if not audio_url:
#             msg = "Failed to upload processed audio to S3."
#             logger.error(msg)
#             raise HTTPException(status_code=500, detail=msg)

#         # 7) Create Audiobook record
#         audiobook_start_time = time.time()
#         duration_sec = len(combined_audio) / 1000
#         hours = int(duration_sec // 3600)
#         minutes = int((duration_sec % 3600) // 60)
#         seconds = int(duration_sec % 60)

#         audiobook_data = AudioBookCreate(
#             title=document.title,
#             author=document.author,
#             narrator="Generated Narrator",
#             duration=f"{hours} hours {minutes} minutes {seconds} seconds",
#             genre=document.genre or "Unknown",
#             publication_date=datetime.utcnow(),
#             file_key=processed_audio_key,
#             is_dolby_atmos_supported=is_dolby_atmos,
#             url=audio_url,
#             document_id=document.id,
#         )
#         audiobook = create_audiobook(db, audiobook_data)
#         document.audiobook = audiobook
#         document.audio_url = audio_url
#         document.processing_status = "completed"
#         db.commit()
#         db.refresh(document)
#         audiobook_end_time = time.time()
#         logger.info(f"[SYNC] Document {document.id} fully processed. Returning result.")
#         # Total processing time
#         total_end_time = time.time()
#         total_processing_time = total_end_time - total_start_time
#         logger.info(f"[SYNC] Total processing time for Document {document.id}: {total_processing_time:.2f} seconds.")
#         return document

#     except HTTPException:
#         raise
#     except Exception as e:
#         logger.error(f"Unhandled error in create_and_process_document: {str(e)}")
#         db.rollback()
#         raise HTTPException(
#             status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
#             detail=str(e),
#         )

# # ------------------------------------------------------------------------------
# #  OTHER DOCUMENT FUNCTIONS
# # ------------------------------------------------------------------------------
# from app.schemas.document_filter import DocumentFilter  # Make sure you have this import

# def get_documents(
#     db: Session,
#     user_id: int,
#     filter_params: DocumentFilter,
#     skip: int = 0,
#     limit: int = 10
# ) -> List[Document]:
#     """
#     Retrieve documents belonging to a user with optional filters (search, status, file_type, date range).
#     """
#     query = db.query(Document).filter(Document.user_id == user_id)

#     # Apply search filter
#     if filter_params.search:
#         search_term = f"%{filter_params.search}%"
#         query = query.filter(
#             or_(Document.title.ilike(search_term), Document.description.ilike(search_term))
#         )
#     # Apply status filter
#     if filter_params.status:
#         query = query.filter(Document.processing_status == filter_params.status)
#     # Apply file type filter
#     if filter_params.file_type:
#         query = query.filter(Document.file_type == filter_params.file_type)
#     # Apply date range filter
#     if filter_params.start_date:
#         query = query.filter(Document.created_at >= filter_params.start_date)
#     if filter_params.end_date:
#         query = query.filter(Document.created_at <= filter_params.end_date)

#     # Apply pagination
#     documents = query.order_by(Document.created_at.desc()).offset(skip).limit(limit).all()
#     return documents


# def count_user_pdfs(db: Session, user_id: int) -> int:
#     """
#     Count how many PDF documents a user has uploaded.
#     Typically used to limit free users to a certain number of PDF uploads.
#     """
#     return db.query(Document).filter(
#         Document.user_id == user_id,
#         Document.file_type == "application/pdf"
#     ).count()


# def update_document(
#     db: Session,
#     document_id: int,
#     user_id: int,
#     document_update: DocumentUpdate
# ) -> Document:
#     """
#     Update an existing document record (title, description, tags, etc.).
#     By default, no re-processing is done here.
#     """
#     try:
#         document = db.query(Document).filter(
#             Document.id == document_id,
#             Document.user_id == user_id
#         ).first()

#         if not document:
#             raise HTTPException(
#                 status_code=status.HTTP_404_NOT_FOUND,
#                 detail="Document not found"
#             )

#         # Update the document fields
#         for key, value in document_update.dict(exclude_unset=True).items():
#             setattr(document, key, value)

#         db.commit()
#         db.refresh(document)
#         return document
#     except Exception as e:
#         db.rollback()
#         logger.error(f"Error updating document {document_id}: {str(e)}")
#         raise HTTPException(
#             status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
#             detail="Failed to update document"
#         )


# def delete_document(db: Session, document_id: int, user_id: int) -> bool:
#     """
#     Delete a document and associated files (original + audio) from S3,
#     then remove the DB record.
#     """
#     try:
#         document = db.query(Document).filter(
#             Document.id == document_id,
#             Document.user_id == user_id
#         ).first()

#         if not document:
#             raise HTTPException(
#                 status_code=status.HTTP_404_NOT_FOUND,
#                 detail="Document not found"
#             )

#         # Delete files from S3
#         if document.file_key:
#             s3_handler.delete_file(settings.S3_BUCKET_NAME, document.file_key)
#         if document.audio_key:
#             s3_handler.delete_file(settings.S3_BUCKET_NAME, document.audio_key)

#         # Delete the DB record
#         db.delete(document)
#         db.commit()
#         return True
#     except Exception as e:
#         db.rollback()
#         logger.error(f"Error deleting document {document_id}: {str(e)}")
#         raise HTTPException(
#             status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
#             detail="Failed to delete document"
#         )

# # ------------------------------------------------------------------------------
# #  OPTIONAL: Search with Redis Caching (if used)
# # ------------------------------------------------------------------------------
# # from app.utils.redis_utils import redis_client
# # import json

# # def search_documents_with_cache(
# #     db: Session,
# #     prefix: str,
# #     bucketname: str,
# #     limits: int,
# #     levels: int,
# #     offsets: int,
# #     user_id: int
# # ):
# #     """
# #     Example: Searching for documents with Redis caching logic (custom).
# #     If you don't need advanced caching, you can remove this.
# #     """
# #     cache_key = f"search:{prefix}:{bucketname}:{limits}:{levels}:{offsets}:{user_id}"
# #     cached_result = redis_client.get(cache_key)

# #     if cached_result:
# #         return json.loads(cached_result)

# #     # ... your custom search logic here ...
# #     # result = db.execute(...).fetchall()

# #     # redis_client.set(cache_key, json.dumps(serialized_result), ex=80640)
# #     # return serialized_result

# #     # Return empty or real data
# #     return []



# app/services/document_service.py
from app.core.config import settings
import io
import json
import logging
import time
from datetime import datetime
from typing import List, Optional, Tuple
from concurrent.futures import ThreadPoolExecutor
from functools import partial

from sqlalchemy.orm import Session, joinedload
from fastapi import HTTPException, status
from docx import Document as DocxDocument
from app.database import SessionLocal
from app.models import Document, Audiobook, User
from app.schemas.document import DocumentCreate, DocumentFilter, DocumentUpdate
from app.schemas.audiobook import AudioBookCreate
from app.services.audiobook_service import create_audiobook
from app.services.ocr_service import OCRService
from app.services.tts_service import TTSService
from app.services.rekognition_service import RekognitionService
from app.utils.s3_utils import s3_handler
from app.utils.lang_utils import map_language_code_to_supported, infer_genre
from app.utils.redis_utils import redis_client, delete_pattern
from pydub import AudioSegment
from pydub.effects import normalize
from pydub.generators import WhiteNoise

logger = logging.getLogger(__name__)

# Initialize external services
ocr_service = OCRService(region_name='us-east-1')
tts_service = TTSService(region_name='us-east-1')
rekognition_service = RekognitionService(region_name='us-east-1')


# ------------------------------------------------------------------------------
#  IMMERSIVE AUDIO EFFECTS
# ------------------------------------------------------------------------------
def apply_immersive_effects(audio: AudioSegment, apply: bool = False) -> AudioSegment:
    """
    Apply immersive audio effects to enhance the listening experience.

    Args:
        audio (AudioSegment): The main audio segment.
        apply (bool): Flag to determine whether to apply immersive effects.

    Returns:
        AudioSegment: The processed audio with or without immersive effects.
    """
    try:
        # Normalize audio volume
        audio = normalize(audio)

        if not apply:
            return audio

        # Apply stereo widening by panning duplicates
        left = audio.pan(-0.3)
        right = audio.pan(0.3)
        immersive_audio = left.overlay(right)

        # Add subtle reverberation to simulate space
        reverberated = immersive_audio.fade_in(500).fade_out(500)

        # Optionally, add white noise to simulate ambient sound
        noise = WhiteNoise().to_audio_segment(duration=len(reverberated)).apply_gain(-30)
        immersive_audio = reverberated.overlay(noise)

        return immersive_audio

    except Exception as e:
        logger.error(f"Error applying immersive effects: {str(e)}")
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail="Failed to apply immersive audio effects."
        )


# ------------------------------------------------------------------------------
#  DATABASE OPERATIONS
# ------------------------------------------------------------------------------

def update_document(
    db: Session,
    document_id: int,
    user_id: int,
    document_update: DocumentUpdate
) -> Document:
    """Update an existing document record with Redis cache invalidation."""
    try:
        document = db.query(Document).filter(
            Document.id == document_id,
            Document.user_id == user_id
        ).first()

        if not document:
            raise HTTPException(
                status_code=status.HTTP_404_NOT_FOUND,
                detail="Document not found"
            )

        # Update the document fields
        for key, value in document_update.dict(exclude_unset=True).items():
            setattr(document, key, value)

        db.commit()
        db.refresh(document)

        # Invalidate relevant Redis cache
        delete_pattern("search:*")

        return document
    except HTTPException:
        raise
    except Exception as e:
        db.rollback()
        logger.error(f"Error updating document {document_id}: {str(e)}")
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail="Failed to update document"
        )


def create_document(
    db: Session,
    document_data: DocumentCreate,
    user_id: int,
    file_key: str,
    file_size: Optional[int] = None
) -> Document:
    """Create a new document record with Redis cache invalidation."""
    try:
        document = Document(
            title=document_data.title,
            author=document_data.author,
            file_type=document_data.file_type,
            file_key=file_key,
            url=document_data.url,
            is_public=document_data.is_public,
            user_id=user_id,
            file_size=file_size,
            upload_date=datetime.utcnow(),
            processing_status="pending",
        )
        db.add(document)
        db.commit()
        db.refresh(document)

        # Invalidate relevant Redis cache
        delete_pattern("search:*")

        return document
    except Exception as e:
        db.rollback()
        logger.error(f"Error creating document: {str(e)}")
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail="Failed to create document"
        )


def get_documents(
    db: Session,
    user_id: int,
    filter_params: DocumentFilter,
    skip: int = 0,
    limit: int = 10
) -> List[Document]:
    """Get documents with filtering."""
    try:
        query = db.query(Document).filter(Document.user_id == user_id)

        if filter_params.search:
            query = query.filter(Document.title.ilike(f"%{filter_params.search}%"))

        if filter_params.status:
            query = query.filter(Document.processing_status == filter_params.status)

        if filter_params.file_type:
            query = query.filter(Document.file_type == filter_params.file_type)

        if filter_params.start_date:
            query = query.filter(Document.upload_date >= filter_params.start_date)

        if filter_params.end_date:
            query = query.filter(Document.upload_date <= filter_params.end_date)

        return query.offset(skip).limit(limit).all()
    except Exception as e:
        logger.error(f"Error retrieving documents: {str(e)}")
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail="Failed to retrieve documents"
        )


def delete_document(db: Session, document_id: int, user_id: int) -> bool:
    """Delete document and associated files with Redis cache invalidation."""
    try:
        document = db.query(Document).filter(
            Document.id == document_id,
            Document.user_id == user_id
        ).first()

        if not document:
            raise HTTPException(
                status_code=status.HTTP_404_NOT_FOUND,
                detail="Document not found"
            )

        # Delete files from S3
        if document.file_key:
            s3_handler.delete_file(settings.S3_BUCKET_NAME, document.file_key)
        if document.audio_key:
            s3_handler.delete_file(settings.S3_BUCKET_NAME, document.audio_key)

        # Delete database record
        db.delete(document)
        db.commit()

        # Invalidate relevant Redis cache
        delete_pattern("search:*")

        return True
    except HTTPException:
        raise
    except Exception as e:
        db.rollback()
        logger.error(f"Error deleting document {document_id}: {str(e)}")
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail="Failed to delete document"
        )


def search_documents_with_cache(db: Session, prefix: str, bucketname: str, limits: int, levels: int, offsets: int, user_id: int):
    """Search documents with Redis caching."""
    try:
        cache_key = f"search:{prefix}:{bucketname}:{limits}:{levels}:{offsets}"
        cached_result = redis_client.get(cache_key)

        if cached_result:
            return json.loads(cached_result)

        query = """
        WITH files_folders AS (
            SELECT ((string_to_array(objects.name, '/'))[%(levels)s]) AS folder
            FROM storage.objects
            WHERE objects.name ILIKE %(prefix)s || '%'
            AND bucket_id = %(bucketname)s
            AND user_id = %(user_id)s
            GROUP BY folder
            LIMIT %(limits)s
            OFFSET %(offsets)s
        )
        SELECT files_folders.folder AS name, objects.id, objects.updated_at,
               objects.created_at, objects.last_accessed_at, objects.metadata
        FROM files_folders
        LEFT JOIN storage.objects
        ON %(prefix)s || files_folders.folder = objects.name
           AND objects.bucket_id = %(bucketname)s
           AND objects.user_id = %(user_id)s;
        """
        result = db.execute(query, {
            "prefix": prefix,
            "bucketname": bucketname,
            "user_id": user_id,
            "limits": limits,
            "levels": levels,
            "offsets": offsets,
        }).fetchall()

        serialized_result = [
            {
                "name": row.name,
                "id": row.id,
                "updated_at": row.updated_at.isoformat(),
                "created_at": row.created_at.isoformat(),
                "last_accessed_at": row.last_accessed_at.isoformat(),
                "metadata": row.metadata,
            }
            for row in result
        ]

        # Cache the result in Redis for future requests
        redis_client.set(cache_key, json.dumps(serialized_result), ex=80640)
        return serialized_result
    except Exception as e:
        logger.error(f"Error searching documents with cache: {str(e)}")
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail="Failed to search documents."
        )


# ------------------------------------------------------------------------------
#  AUDIO PROCESSING
# ------------------------------------------------------------------------------

def process_text_chunks_in_parallel(text_chunks: List[str], detected_language: str, max_workers: int = 8) -> List[bytes]:
    """Process text chunks for TTS in parallel."""
    try:
        with ThreadPoolExecutor(max_workers=max_workers) as executor:
            logger.info(f"Processing {len(text_chunks)} text chunks for TTS.")
            audio_chunks = list(executor.map(
                lambda chunk: tts_service.convert_text_to_speech(chunk, detected_language),
                text_chunks
            ))
        return audio_chunks
    except Exception as e:
        logger.error(f"Error processing text chunks in parallel: {str(e)}")
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail="Failed to process text chunks for TTS."
        )


def generate_audio_for_document(document: Document, text: str, detected_language: str, apply_immersive: bool = False) -> Tuple[str, str]:
    """Generate audio for a document and upload to S3."""
    try:
        # Split text into manageable chunks
        max_length = 3000
        text_chunks = [text[i:i + max_length] for i in range(0, len(text), max_length)]

        # Process chunks in parallel
        audio_chunks = process_text_chunks_in_parallel(text_chunks, detected_language)

        # Combine audio chunks into one audio file
        audio_combined = AudioSegment.empty()
        for chunk_audio in audio_chunks:
            audio_segment = AudioSegment.from_file(io.BytesIO(chunk_audio), format="mp3")
            audio_combined += audio_segment

        # Apply immersive audio effects
        immersive_audio = apply_immersive_effects(audio_combined, apply=apply_immersive)

        # Export processed audio to a BytesIO object
        audio_bytes_io = io.BytesIO()
        immersive_audio.export(audio_bytes_io, format="mp3")
        audio_bytes_io.seek(0)

        # Define S3 key for the processed audio
        processed_audio_key = f"{settings.S3_FOLDER_NAME}/audio/{document.user_id}/{document.id}_processed.mp3"
        logger.info(f"Uploading processed audio to S3 with key: {processed_audio_key}.")

        # Upload processed audio to S3
        audio_url = s3_handler.upload_file(audio_bytes_io, settings.S3_BUCKET_NAME, processed_audio_key, "audio/mpeg")
        logger.info(f"Processed audio uploaded to S3 successfully. URL: {audio_url}.")

        return audio_url, processed_audio_key
    except Exception as e:
        logger.error(f"Error generating audio for document {document.id}: {str(e)}")
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail="Failed to generate audio for the document."
        )


# ------------------------------------------------------------------------------
#  DOCUMENT PROCESSING
# ------------------------------------------------------------------------------

def create_and_process_document(
    db: Session,
    document_data: DocumentCreate,
    user: User,
    file_content: bytes,
    subscription_type: str = "free"
) -> Document:
    """
    Create a new Document, perform OCR + TTS + S3 uploads **synchronously**,
    and attach the audiobook to the Document record. Return the final Document.
    """
    try:
        # Start total processing time
        total_start_time = time.time()
        
        # 1) Insert Document row in DB (initially "processing")
        document = Document(
            title=document_data.title,
            author=document_data.author,
            file_type=document_data.file_type,
            file_key=document_data.file_key,
            url=document_data.url,
            is_public=document_data.is_public,
            user_id=user.id,
            file_size=len(file_content),
            upload_date=datetime.utcnow(),
            created_at=document_data.created_at,
            processing_status='processing'
        )
        db.add(document)
        db.commit()
        db.refresh(document)

        logger.info(f"[SYNC] Document record {document.id} created. Starting processing...")

        # 2) OCR / Text Extraction
        ocr_start_time = time.time()
        if document.file_type == "application/pdf":
            text, detected_lang = ocr_service.extract_text_from_pdf(
                bucket_name=settings.S3_BUCKET_NAME,
                file_key=document.file_key
            )

        elif document.file_type.startswith("image/"):
            text, detected_lang = ocr_service.extract_text_from_image(
                bucket_name=settings.S3_BUCKET_NAME,
                file_key=document.file_key
            )
            # Rekognition tags for images
            tags_start_time = time.time()
            tags = rekognition_service.detect_labels(settings.S3_BUCKET_NAME, document.file_key)
            if isinstance(tags, str):
                tags = tags.split(", ")
            document.tags = tags or ["No tags detected"]
            tags_end_time = time.time()
            logger.info(f"[SYNC] Rekognition tags for Document {document.id} processed in {tags_end_time - tags_start_time:.2f} seconds.")

        elif document.file_type == "text/plain":
            # Read text directly from S3
            file_bytes = s3_handler.get_file_content(settings.S3_BUCKET_NAME, document.file_key)
            text = file_bytes.decode("utf-8")
            detected_lang = 'en'

        elif document.file_type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
            # Handle DOCX
            file_bytes = s3_handler.get_file_content(settings.S3_BUCKET_NAME, document.file_key)
            with io.BytesIO(file_bytes) as docx_stream:
                docx_obj = DocxDocument(docx_stream)
                text = "\n".join(p.text for p in docx_obj.paragraphs)
            detected_lang = ocr_service.detect_language(text)

        else:
            msg = f"Unsupported file type: {document.file_type}"
            logger.error(msg)
            raise HTTPException(status_code=400, detail=msg)

        if not text.strip():
            msg = f"No text found in document {document.id}"
            logger.error(msg)
            raise HTTPException(status_code=400, detail=msg)
        
        ocr_end_time = time.time()
        logger.info(f"[SYNC] OCR for Document {document.id} completed in {ocr_end_time - ocr_start_time:.2f} seconds.")

        # 3) Post-OCR updates
        post_ocr_start_time = time.time()
        document.detected_language = map_language_code_to_supported(detected_lang or "en")
        document.description = text[:200].strip()  # First 200 chars
        document.genre = infer_genre(text)
        db.commit()
        db.refresh(document)
        post_ocr_end_time = time.time()
        logger.info(f"[SYNC] Post-OCR updates for Document {document.id} completed in {post_ocr_end_time - post_ocr_start_time:.2f} seconds.")

        # 4) TTS Generation
        tts_start_time = time.time()
        is_dolby_atmos = (subscription_type == "premium")
        logger.info(f"[SYNC] Generating audio for Document {document.id} (Dolby={is_dolby_atmos})")

        # Split large text into chunks of ~3000 chars
        max_chars = 3000
        if len(text) <= max_chars:
            text_chunks = [text]
        else:
            text_chunks = [text[i : i + max_chars] for i in range(0, len(text), max_chars)]

        # Determine the number of workers based on the number of chunks and system resources
        max_workers = min(12, len(text_chunks))  # Adjust as needed

        # Process chunks in parallel using the predefined function
        polly_start_time = time.time()
        polly_results = process_text_chunks_in_parallel(text_chunks, detected_lang or "en", max_workers=max_workers)
        polly_end_time = time.time()
        logger.info(f"[SYNC] TTS conversion for Document {document.id} completed in {polly_end_time - polly_start_time:.2f} seconds.")

        # Combine all segments
        audio_segments = []
        for chunk_data in polly_results:
            if chunk_data:
                audio_segment = AudioSegment.from_mp3(io.BytesIO(chunk_data))
                audio_segments.append(audio_segment)

        if not audio_segments:
            msg = f"TTS failed: no audio segments produced for document {document.id}"
            logger.error(msg)
            raise HTTPException(status_code=500, detail=msg)

        combined_audio = sum(audio_segments)

        # Apply immersive effects for premium users
        if is_dolby_atmos:
            immersive_start_time = time.time()
            combined_audio = apply_immersive_effects(combined_audio, apply=True)
            immersive_end_time = time.time()
            logger.info(f"[SYNC] Applied immersive effects to Document {document.id} in {immersive_end_time - immersive_start_time:.2f} seconds.")

        # 5) Export final MP3
        export_start_time = time.time()
        final_audio_bytes = io.BytesIO()
        combined_audio.export(final_audio_bytes, format="mp3", bitrate="128k")
        final_audio_bytes.seek(0)
        export_end_time = time.time()
        logger.info(f"[SYNC] Exported audio for Document {document.id} in {export_end_time - export_start_time:.2f} seconds.")

        # 6) Upload processed MP3 to S3
        upload_start_time = time.time()
        processed_audio_key = f"audio/{user.id}/{document.id}_processed.mp3"
        audio_url = s3_handler.upload_file(
            file_obj=final_audio_bytes,
            bucket=settings.S3_BUCKET_NAME,
            key=processed_audio_key,
            content_type="audio/mpeg"
        )
        upload_end_time = time.time()
        if not audio_url:
            msg = "Failed to upload processed audio to S3."
            logger.error(msg)
            raise HTTPException(status_code=500, detail=msg)

        # 7) Create Audiobook record
        audiobook_start_time = time.time()
        duration_sec = len(combined_audio) / 1000
        hours = int(duration_sec // 3600)
        minutes = int((duration_sec % 3600) // 60)
        seconds = int(duration_sec % 60)

        audiobook_data = AudioBookCreate(
            title=document.title,
            author=document.author,
            narrator="Generated Narrator",
            duration=f"{hours} hours {minutes} minutes {seconds} seconds",
            genre=document.genre or "Unknown",
            publication_date=datetime.utcnow(),
            file_key=processed_audio_key,
            is_dolby_atmos_supported=is_dolby_atmos,
            url=audio_url,
            document_id=document.id,
        )
        audiobook = create_audiobook(db, audiobook_data)
        document.audiobook = audiobook
        document.audio_url = audio_url
        document.processing_status = "completed"
        db.commit()
        db.refresh(document)
        audiobook_end_time = time.time()
        logger.info(f"[SYNC] Audiobook created with ID {audiobook.id} for Document ID {document.id}.")

        # 8) Eagerly Load Relationships Before Returning
        # Use joinedload to ensure 'audiobook' is loaded
        document = db.query(Document).options(joinedload(Document.audiobook)).filter(Document.id == document.id).first()

        logger.info(f"[SYNC] Document {document.id} fully processed. Returning result.")

        # Total processing time
        total_end_time = time.time()
        total_processing_time = total_end_time - total_start_time
        logger.info(f"[SYNC] Total processing time for Document {document.id}: {total_processing_time:.2f} seconds.")
        return document

    except HTTPException as http_exc:
        logger.error(f"HTTP error in create_and_process_document: {str(http_exc.detail)}")
        raise http_exc
    except Exception as e:
        logger.error(f"Unhandled error in create_and_process_document: {str(e)}")
        db.rollback()
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail=str(e),
        )
    finally:
        db.close()



def count_user_pdfs(db: Session, user_id: int) -> int:
    """
    Count how many PDF documents a user has uploaded.
    Typically used to limit free users to a certain number of PDF uploads.
    """
    return db.query(Document).filter(
        Document.user_id == user_id,
        Document.file_type == "application/pdf"
    ).count()
